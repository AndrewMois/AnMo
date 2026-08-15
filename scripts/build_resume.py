from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "documents" / "Andriy-Moiseyenko-Resume.docx"

INK = RGBColor(0x10, 0x20, 0x27)
GRAPHITE = RGBColor(0x17, 0x19, 0x18)
GOLD = RGBColor(0x76, 0x53, 0x1E)
STEEL = RGBColor(0x50, 0x6A, 0x78)
HAIRLINE = "D8D1C5"
FONT = "Arial"


def set_font(run, size, color=GRAPHITE, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def set_spacing(paragraph, before=0, after=0, line=1.0, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.widow_control = True


def add_bottom_border(paragraph, color=HAIRLINE, size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="Resume Section")
    set_spacing(paragraph, before=5, after=3, line=1.0, keep_next=True)
    run = set_font(paragraph.add_run(text.upper()), 8.4, GOLD, bold=True)
    run._element.get_or_add_rPr().append(OxmlElement("w:spacing"))
    run._element.rPr[-1].set(qn("w:val"), "24")
    add_bottom_border(paragraph)
    return paragraph


def add_labeled_line(doc, label, text):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, after=1.2, line=1.0)
    set_font(paragraph.add_run(f"{label}: "), 8.35, INK, bold=True)
    set_font(paragraph.add_run(text), 8.35, GRAPHITE)
    return paragraph


def add_job(doc, title, dates, organization, location, bullets):
    paragraph = doc.add_paragraph(style="Resume Job")
    set_spacing(paragraph, before=2.5, after=0.5, line=1.0, keep_next=True)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.18), WD_TAB_ALIGNMENT.RIGHT)
    set_font(paragraph.add_run(title.upper()), 8.8, INK, bold=True)
    set_font(paragraph.add_run(f"\t{dates.upper()}"), 8.1, GOLD, bold=True)

    organization_line = doc.add_paragraph()
    set_spacing(organization_line, after=1.5, line=1.0, keep_next=True)
    set_font(organization_line.add_run(organization), 8.45, GRAPHITE, bold=True)
    set_font(organization_line.add_run(f"  |  {location}"), 8.2, STEEL)

    for text in bullets:
        bullet = doc.add_paragraph(style="Resume Bullet")
        set_spacing(bullet, after=1.35, line=1.02)
        bullet.paragraph_format.left_indent = Inches(0.22)
        bullet.paragraph_format.first_line_indent = Inches(-0.14)
        set_font(bullet.add_run(text), 8.25, GRAPHITE)


def add_education(doc, credential, dates, institution, detail=""):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, before=1.5, after=0.3, line=1.0, keep_next=True)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.18), WD_TAB_ALIGNMENT.RIGHT)
    set_font(paragraph.add_run(credential), 8.5, INK, bold=True)
    set_font(paragraph.add_run(f"\t{dates}"), 7.9, GOLD, bold=True)
    line = doc.add_paragraph()
    set_spacing(line, after=1.2, line=1.0)
    set_font(line.add_run(institution), 8.15, GRAPHITE)
    if detail:
        set_font(line.add_run(f"  |  {detail}"), 8.0, STEEL)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(8.35)
    normal.font.color.rgb = GRAPHITE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for name in ["Resume Section", "Resume Job"]:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal

    bullet = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(8.25)
    bullet.font.color.rgb = GRAPHITE


def build_resume():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.46)
    section.left_margin = Inches(0.66)
    section.right_margin = Inches(0.66)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    doc.core_properties.title = "Andriy Moiseyenko — QA Automation Engineer"
    doc.core_properties.subject = "Professional resume"
    doc.core_properties.author = "Andriy Moiseyenko"
    doc.core_properties.keywords = "QA automation, quality engineering, SDET, software testing"

    name = doc.add_paragraph()
    set_spacing(name, after=0, line=0.95, keep_next=True)
    set_font(name.add_run("ANDRIY MOISEYENKO"), 24, INK, bold=True)

    role = doc.add_paragraph()
    set_spacing(role, after=3, line=1.0, keep_next=True)
    role_run = set_font(role.add_run("QA AUTOMATION ENGINEER  ·  QUALITY SYSTEMS"), 9.6, GOLD, bold=True)
    role_run._element.get_or_add_rPr().append(OxmlElement("w:spacing"))
    role_run._element.rPr[-1].set(qn("w:val"), "18")

    contact = doc.add_paragraph()
    set_spacing(contact, after=5, line=1.0, keep_next=True)
    set_font(
        contact.add_run(
            "Ottawa, ON  ·  613-407-7125  ·  amoiseyenko@ukr.net  ·  anmo.dev  ·  linkedin.com/in/mrskyray  ·  github.com/AndrewMois"
        ),
        7.9,
        STEEL,
    )
    add_bottom_border(contact, color="B48A3C", size="8", space="5")

    add_section_heading(doc, "Profile")
    profile = doc.add_paragraph()
    set_spacing(profile, after=1.5, line=1.04)
    set_font(
        profile.add_run(
            "QA Automation Engineer who builds quality systems from first principles across interfaces, APIs, backend services, data workflows, performance, and AI-driven behaviour. I combine test strategy with hands-on framework engineering, CI/CD, root-cause analysis, and a developer’s understanding of complete products."
        ),
        8.55,
        GRAPHITE,
    )

    add_section_heading(doc, "Technical strengths")
    add_labeled_line(
        doc,
        "Quality engineering",
        "Test strategy, automation architecture, release readiness, negative testing, root-cause analysis",
    )
    add_labeled_line(
        doc,
        "Automation and services",
        "Java, Selenium, TestNG, WebdriverIO, Appium, Mocha, Bruno, Postman/Newman, SOAP/XML",
    )
    add_labeled_line(
        doc,
        "Performance, delivery, product",
        "k6, Locust, Jenkins, Docker, Git, Jira/Xray, SvelteKit, TypeScript, MongoDB, Sanity",
    )

    add_section_heading(doc, "Professional experience")
    add_job(
        doc,
        "QA Automation Engineer",
        "Mar 2025 – Present",
        "BluWave AI",
        "Ottawa, ON",
        [
            "Own the testing approach for AI-powered energy products across web, Android/iOS, APIs, backend services, data workflows, and forecasting behaviour.",
            "Architected automation foundations from scratch for UI/mobile (WebdriverIO, Appium, Mocha), API regression (Bruno, Postman/Newman), SOAP/XML services, and performance programs (k6, Locust).",
            "Integrated reliable suites into Dockerized Jenkins pipelines with environment resolution, secure configuration, normalized reporting, and Jira/Xray publishing.",
            "Designed risk-based API coverage for identity, access control, CRUD, filtering, persistence, schemas, negative paths, and deterministic cleanup.",
            "Validate forecasting quality through scenario-based testing and measures including MAE, RMSE, MAPE, correlation, bias, peak miss, and lag; diagnose failures across logs, APIs, dashboards, devices, and CI evidence.",
        ],
    )

    add_job(
        doc,
        "QA Automation Engineer",
        "Sep 2023 – Aug 2024",
        "Canada Revenue Agency",
        "Ottawa, ON",
        [
            "Developed Java automation with Selenium and TestNG, translating user stories and acceptance criteria into focused, reusable regression scenarios.",
            "Created a reusable negative-testing framework and integrated automation into Jenkins for repeatable evidence, triage, and developer collaboration.",
            "Led test design and automation structure for a key initiative; reviewed work, mentored junior testers, and coordinated technical decisions with stakeholders.",
        ],
    )

    add_job(
        doc,
        "IT Technician",
        "2021 – 2022",
        "UkrTechProm",
        "Ukraine",
        [
            "Maintained network infrastructure, websites, and internal applications while diagnosing hardware, software, and connectivity issues for end users.",
        ],
    )

    add_section_heading(doc, "Selected product")
    product = doc.add_paragraph()
    set_spacing(product, after=1.2, line=1.02)
    set_font(product.add_run("YOUNG MYSTIC"), 8.45, INK, bold=True)
    set_font(
        product.add_run(
            "  ·  Multilingual mobile-first essential-oil library used by 200+ people; built and maintained with SvelteKit, MongoDB, Sanity, authentication, and carefully bounded PWA/offline behaviour."
        ),
        8.25,
        GRAPHITE,
    )

    add_section_heading(doc, "Education and credentials")
    add_education(
        doc,
        "Web Development & Internet Applications Diploma",
        "Sep 2022 – Aug 2024",
        "Algonquin College · Ottawa, ON",
        "Dean’s Honour List · GPA 3.9 / 4.0",
    )
    add_education(
        doc,
        "Master of Pharmacy",
        "Sep 2016 – Aug 2021",
        "Bogomolets National Medical University · Kyiv, Ukraine",
    )

    additional = doc.add_paragraph()
    set_spacing(additional, before=2.5, after=0, line=1.0)
    set_font(additional.add_run("RELIABILITY STATUS"), 7.8, GOLD, bold=True)
    set_font(additional.add_run("  Valid through August 2033"), 8.0, GRAPHITE)
    set_font(additional.add_run("     LANGUAGES"), 7.8, GOLD, bold=True)
    set_font(additional.add_run("  English, Ukrainian, Russian · French (basic)"), 8.0, GRAPHITE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(footer, before=0, after=0, line=1.0)
    set_font(footer.add_run("anmo.dev"), 7.3, STEEL)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_resume()
